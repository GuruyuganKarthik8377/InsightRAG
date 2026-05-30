"""
Full integration + unit test suite for the RAG pipeline.
Uses real models — no mocks. Tests actual retrieval quality.
Run from the backend/ directory:
    pytest tests/test_pipeline.py -v
"""
import os
import sys
import tempfile

import numpy as np
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from app.ingestion.loader import load_document
from app.ingestion.chunker import chunk_text
from app.ingestion.embedder import embed_docs, embed_query
from app.db.vector_store import VectorStore
from app.retrieval.reranker import rerank
from app.schemas.models import QueryResponse

SAMPLE_TEXT = """
Artificial intelligence is transforming industries worldwide.
Machine learning enables computers to learn from data without explicit programming.
Deep learning uses neural networks with many layers to process complex information.
Natural language processing allows machines to understand and generate human language.
Computer vision enables machines to interpret and understand visual information from images.
""" * 10  # repeat to get enough content for chunking


# ── Unit Tests ───────────────────────────────────────────────────────────────

def test_chunker_output_format():
    pages = [{"text": SAMPLE_TEXT, "page": 1}]
    chunks = chunk_text(pages, doc_id="test-doc", source="test.txt")
    assert len(chunks) > 0
    for c in chunks:
        assert "chunk_id" in c
        assert "text" in c
        assert "doc_id" in c
        assert "source" in c
        assert len(c["text"]) >= 30
        assert len(c["text"]) <= 800


def test_chunker_no_empty_chunks():
    pages = [{"text": "   \n\n   ", "page": 1}]
    chunks = chunk_text(pages, doc_id="x", source="empty.txt")
    assert len(chunks) == 0


def test_chunker_page_none_propagated():
    pages = [{"text": SAMPLE_TEXT, "page": None}]
    chunks = chunk_text(pages, doc_id="d", source="doc.docx")
    assert len(chunks) > 0
    for c in chunks:
        assert c["page"] is None


def test_chunker_page_int_propagated():
    pages = [{"text": SAMPLE_TEXT, "page": 7}]
    chunks = chunk_text(pages, doc_id="d", source="doc.pdf")
    assert len(chunks) > 0
    for c in chunks:
        assert c["page"] == 7


def test_embed_docs_shape():
    texts = ["hello world", "machine learning is great"]
    vecs = embed_docs(texts)
    assert vecs.shape == (2, 1024)
    assert vecs.dtype == np.float32


def test_embed_query_shape():
    vec = embed_query("what is machine learning?")
    assert vec.shape == (1, 1024)
    assert vec.dtype == np.float32


def test_embed_docs_normalized():
    texts = ["deep learning uses neural networks"]
    vecs = embed_docs(texts)
    norm = float(np.linalg.norm(vecs[0]))
    assert abs(norm - 1.0) < 1e-4, f"Expected norm ~1.0, got {norm}"


def test_embed_query_normalized():
    vec = embed_query("what is AI?")
    norm = float(np.linalg.norm(vec[0]))
    assert abs(norm - 1.0) < 1e-4, f"Expected norm ~1.0, got {norm}"


def test_embed_query_prefix_applied():
    text = "machine learning"
    doc_vec = embed_docs([text])
    query_vec = embed_query(text)
    similarity = float(np.dot(doc_vec[0], query_vec[0]))
    # Same text with/without prefix — similar but not identical
    assert 0.7 < similarity < 1.0, f"Similarity out of expected range: {similarity}"


def test_vector_store_add_and_search():
    import faiss

    vs = VectorStore()
    vs.index = faiss.IndexFlatIP(1024)
    vs.metadata = []

    texts = ["deep learning uses neural networks", "python is a programming language"]
    vecs = embed_docs(texts)
    meta = [
        {"chunk_id": "c1", "doc_id": "d1", "text": texts[0], "source": "a.pdf", "page": 1},
        {"chunk_id": "c2", "doc_id": "d1", "text": texts[1], "source": "a.pdf", "page": 2},
    ]
    # Bypass add() to avoid disk write in unit test
    vs.index.add(vecs.astype(np.float32))
    vs.metadata.extend(meta)

    query_vec = embed_query("neural network architecture")
    results = vs.search(query_vec, k=2)
    assert len(results) > 0
    assert results[0]["text"] == texts[0]  # deep learning chunk should rank first


def test_vector_store_empty_search():
    import faiss

    vs = VectorStore()
    vs.index = faiss.IndexFlatIP(1024)
    vs.metadata = []
    query_vec = embed_query("anything")
    results = vs.search(query_vec, k=5)
    assert results == []


def test_vector_store_count_mismatch_raises():
    import faiss
    import numpy as np

    vs = VectorStore()
    vs.index = faiss.IndexFlatIP(1024)
    vs.metadata = []

    vecs = np.random.rand(3, 1024).astype(np.float32)
    meta = [{"chunk_id": "x"}]  # intentional mismatch
    with pytest.raises(ValueError, match="must match"):
        vs.add(vecs, meta)


def test_reranker_output_count():
    candidates = [
        {"chunk_id": f"c{i}", "text": f"sentence number {i} about AI", "score": 0.5}
        for i in range(10)
    ]
    results = rerank("what is AI?", candidates, top_n=5)
    assert len(results) == 5


def test_reranker_scores_in_range():
    candidates = [
        {"chunk_id": "c1", "text": "machine learning trains models on data", "score": 0.9},
        {"chunk_id": "c2", "text": "the sky is blue on a clear day", "score": 0.3},
    ]
    results = rerank("how does machine learning work?", candidates, top_n=2)
    for r in results:
        assert 0.0 <= r["score"] <= 1.0, f"Score out of range: {r['score']}"


def test_reranker_scores_descending():
    candidates = [
        {"chunk_id": f"c{i}", "text": f"result {i}", "score": 0.5}
        for i in range(5)
    ]
    results = rerank("test query", candidates, top_n=5)
    for i in range(len(results) - 1):
        assert results[i]["score"] >= results[i + 1]["score"]


def test_reranker_empty_input():
    results = rerank("test query", [], top_n=5)
    assert results == []


def test_reranker_single_candidate():
    candidates = [{"chunk_id": "c1", "text": "only one result here", "score": 0.5}]
    results = rerank("test", candidates, top_n=5)
    assert len(results) == 1


def test_reranker_overflow_guard():
    """Extreme raw scores must not raise OverflowError."""
    candidates = [
        {"chunk_id": "c1", "text": "extreme score test", "score": 0.0},
    ]
    # Directly patch the compute_score result to simulate extreme values
    import unittest.mock as mock
    import app.retrieval.reranker as reranker_mod

    with mock.patch.object(reranker_mod.reranker, "compute_score", return_value=-1000.0):
        results = rerank("test", candidates, top_n=1)
        assert results[0]["score"] == 0.0  # sigmoid(-20) rounds to 0.0

    with mock.patch.object(reranker_mod.reranker, "compute_score", return_value=1000.0):
        results = rerank("test", candidates, top_n=1)
        assert results[0]["score"] == 1.0  # sigmoid(20) rounds to 1.0


# ── Integration Tests ────────────────────────────────────────────────────────

def test_full_pipeline_txt():
    """End-to-end: text file → chunks → embed → store → retrieve → rerank."""
    import faiss
    from pathlib import Path

    vs = VectorStore()
    vs.index = faiss.IndexFlatIP(1024)
    vs.metadata = []

    with tempfile.NamedTemporaryFile(
        suffix=".txt", mode="w", delete=False, encoding="utf-8"
    ) as f:
        f.write(SAMPLE_TEXT)
        tmp_path = f.name

    try:
        pages = load_document(tmp_path)
        assert len(pages) > 0
        assert pages[0]["page"] is None  # TXT has no page numbers

        chunks = chunk_text(pages, doc_id="int-test", source="test.txt")
        assert len(chunks) > 0

        texts = [c["text"] for c in chunks]
        vecs = embed_docs(texts)
        assert vecs.shape == (len(chunks), 1024)

        # Use internal add to skip disk write
        vs.index.add(vecs.astype(np.float32))
        vs.metadata.extend(chunks)

        query_vec = embed_query("what is deep learning?")
        candidates = vs.search(query_vec, k=20)
        assert len(candidates) > 0

        results = rerank("what is deep learning?", candidates, top_n=5)
        assert len(results) > 0

        top_text = results[0]["text"].lower()
        assert any(kw in top_text for kw in ["deep", "neural", "learning"])
    finally:
        Path(tmp_path).unlink(missing_ok=True)


def test_full_pipeline_pdf(tmp_path):
    """End-to-end: generate a real PDF, load it, assert page numbers are real ints."""
    pytest.importorskip("reportlab", reason="reportlab needed to generate test PDFs")
    from reportlab.pdfgen import canvas

    pdf_path = tmp_path / "test.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "Deep learning uses neural networks with many layers.")
    c.showPage()
    c.drawString(100, 750, "Natural language processing enables machines to understand text.")
    c.save()

    pages = load_document(str(pdf_path))
    assert len(pages) == 2
    assert pages[0]["page"] == 1
    assert pages[1]["page"] == 2
    for p in pages:
        assert isinstance(p["page"], int)
        assert p["text"].strip() != ""


def test_full_pipeline_docx(tmp_path):
    """End-to-end: DOCX loading returns page=None and filters empty paragraphs."""
    import docx

    doc_path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("")  # empty paragraph — should be filtered
    doc.add_paragraph("Machine learning enables computers to learn from data.")
    doc.add_paragraph("   ")  # whitespace-only paragraph — should be filtered
    doc.add_paragraph("Deep learning uses neural networks.")
    doc.save(str(doc_path))

    pages = load_document(str(doc_path))
    assert len(pages) == 1
    assert pages[0]["page"] is None
    assert "" not in pages[0]["text"].split("\n")


def test_persistence_roundtrip(tmp_path):
    """Save index → reload into fresh VectorStore → query → same top result."""
    import faiss

    idx_path = str(tmp_path / "index.faiss")
    meta_path = str(tmp_path / "meta.pkl")

    vs = VectorStore(index_path=idx_path, meta_path=meta_path)

    texts = ["neural networks learn from data", "gradient descent optimizes weights"]
    vecs = embed_docs(texts)
    meta = [
        {"chunk_id": "p1", "text": texts[0], "source": "x.pdf", "page": 1, "doc_id": "d"},
        {"chunk_id": "p2", "text": texts[1], "source": "x.pdf", "page": 2, "doc_id": "d"},
    ]
    vs.add(vecs, meta)
    assert vs.index.ntotal == 2

    vs2 = VectorStore(index_path=idx_path, meta_path=meta_path)
    vs2.load()
    assert vs2.index.ntotal == 2
    assert len(vs2.metadata) == 2

    query_vec = embed_query("how do neural networks learn?")
    results = vs2.search(query_vec, k=2)
    assert len(results) == 2
    assert results[0]["chunk_id"] == "p1"


def test_persistence_desync_raises(tmp_path):
    """Writing extra vector without updating metadata must raise RuntimeError on load."""
    import faiss
    import pickle

    idx_path = str(tmp_path / "index.faiss")
    meta_path = str(tmp_path / "meta.pkl")

    index = faiss.IndexFlatIP(1024)
    vecs = np.random.rand(3, 1024).astype(np.float32)
    index.add(vecs)
    faiss.write_index(index, idx_path)

    # Write only 2 metadata entries — intentional desync
    with open(meta_path, "wb") as f:
        pickle.dump([{"chunk_id": "c1"}, {"chunk_id": "c2"}], f)

    vs = VectorStore(index_path=idx_path, meta_path=meta_path)
    with pytest.raises(RuntimeError, match="desync"):
        vs.load()


def test_txt_latin1_fallback(tmp_path):
    """TXT files encoded in Latin-1 must load without UnicodeDecodeError."""
    txt_path = tmp_path / "latin1.txt"
    latin1_text = "Caf\xe9 and na\xefve approaches to AI.\nMore Latin-1 content here."
    txt_path.write_bytes(latin1_text.encode("latin-1"))

    pages = load_document(str(txt_path))
    assert len(pages) == 1
    assert "Caf" in pages[0]["text"]
    assert pages[0]["page"] is None


def test_query_response_schema():
    """QueryResponse must serialise to valid JSON with all required fields."""
    import json
    from app.schemas.models import RetrievalResult

    result = RetrievalResult(
        text="test text",
        score=0.85,
        source="test.pdf",
        page=3,
        chunk_id="abc-123",
    )
    response = QueryResponse(query="test", results=[result])
    serialized = json.loads(response.model_dump_json())
    assert serialized["query"] == "test"
    assert serialized["results"][0]["score"] == 0.85
    assert serialized["results"][0]["page"] == 3
    assert serialized["message"] is None


def test_query_response_schema_null_page():
    """page=None must serialise as null (not raise)."""
    import json
    from app.schemas.models import RetrievalResult

    result = RetrievalResult(
        text="from a docx file",
        score=0.72,
        source="notes.docx",
        page=None,
        chunk_id="xyz-456",
    )
    response = QueryResponse(query="test", results=[result], message="ok")
    serialized = json.loads(response.model_dump_json())
    assert serialized["results"][0]["page"] is None
    assert serialized["message"] == "ok"
